from flask import Flask, request, send_file, send_from_directory
from flask_cors import CORS
from docx import Document
from docx.shared import Pt, Cm, Inches
from PIL import Image
import io

app = Flask(__name__)
CORS(app)

@app.route('/')
def index():
    return send_from_directory('.', 'index.html')

@app.route('/create-doc', methods=['POST'])
def create_doc():
    files = sorted(request.files.getlist('files'), key=lambda f: f.filename)
    title = request.form.get('title')

    doc = Document()

    # Set page size to A4
    section = doc.sections[-1]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    # Clear header
    for section in doc.sections:
        header = section.header
        if header is not None:
            for paragraph in header.paragraphs:
                for run in paragraph.runs:
                    run.clear()

    # Add custom title
    title_paragraph = doc.add_paragraph(title)
    title_paragraph.alignment = 1  # center alignment
    title_paragraph.runs[0].bold = True
    title_paragraph.runs[0].font.size = Pt(20)

    # Create a table to arrange the images (4 images per row, 4 rows per page)
    table = doc.add_table(rows=1, cols=4)
    table.autofit = True
    row = table.rows[0]

    for i, file in enumerate(files):
        # Determine image orientation
        image = Image.open(file)
        width, height = image.size

        cell = row.cells[i % 4]
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()

        if width > height:
            # Landscape: Set height to 4.08 cm
            run.add_picture(file, height=Cm(4.08))
        else:
            # Portrait: Set width to 4.08 cm
            run.add_picture(file, width=Cm(4.08))

        # Create a new row after every fourth image
        if (i + 1) % 4 == 0:
            row = table.add_row()

        # Create a new table after every sixteenth image (to start a new page)
        if (i + 1) % 16 == 0:
            doc.add_page_break()
            table = doc.add_table(rows=1, cols=4)
            row = table.rows[0]

    # Save document to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Convert title to a valid filename
    valid_filename = "".join([c if c.isalnum() else "_" for c in title]) + ".docx"

    return send_file(buffer, as_attachment=True, download_name=valid_filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

if __name__ == '__main__':
    app.run(debug=True)
